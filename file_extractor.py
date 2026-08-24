# -*- coding: utf-8 -*-
"""
첨부파일 텍스트 추출 모듈

나라장터 공고에 붙는 첨부파일(제안요청서, 규격서 등)은 보통 HWP/HWPX/PDF/DOCX/ZIP 형태입니다.
여기서는 "면접" 같은 특정 단어가 본문에 있는지만 확인하면 되므로,
완벽한 서식 보존보다는 "최대한 많은 텍스트를 뽑아내는 것"을 목표로 합니다.

지원 형식: .hwp, .hwpx, .pdf, .docx, .xlsx, .zip (zip 안에 위 형식이 있으면 재귀적으로 처리)
"""
import io
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

import requests

try:
    import olefile
    HAS_OLEFILE = True
except ImportError:
    HAS_OLEFILE = False

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def extract_from_hwp(file_bytes: bytes) -> str:
    """
    HWP(구 버전, OLE 구조) 파일에서 텍스트 추출.
    완벽하진 않지만 PrvText(미리보기 텍스트)와 BodyText 스트림에서
    한글/영문/숫자 위주로 텍스트를 긁어옵니다. 키워드 검색 용도로는 충분합니다.
    """
    if not HAS_OLEFILE:
        return ""
    try:
        ole = olefile.OleFileIO(io.BytesIO(file_bytes))
        if ole.exists("EncryptedPackage"):
            return ""  # 암호화된 파일은 추출 불가

        text_parts = []
        if ole.exists("PrvText"):
            try:
                raw = ole.openstream("PrvText").read()
                decoded = raw.decode("utf-16-le", errors="ignore").replace("\x00", "")
                if decoded.strip():
                    text_parts.append(decoded)
            except Exception:
                pass

        i = 0
        while ole.exists(f"BodyText/Section{i}"):
            try:
                raw = ole.openstream(f"BodyText/Section{i}").read()
                decoded = raw.decode("utf-16-le", errors="ignore")
                readable = re.findall(r"[\uAC00-\uD7A3a-zA-Z0-9\s.,!?()\[\]/-]+", decoded)
                text_parts.extend(readable)
            except Exception:
                pass
            i += 1

        ole.close()
        return "\n".join(text_parts)
    except Exception:
        return ""


def extract_from_hwpx(file_bytes: bytes) -> str:
    """HWPX(신 버전, ZIP+XML 구조) 파일에서 텍스트 추출."""
    try:
        text_parts = []
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            section_files = sorted(
                n for n in zf.namelist()
                if n.startswith("Contents/section") and n.endswith(".xml")
            )
            for name in section_files:
                try:
                    content = zf.read(name).decode("utf-8", errors="ignore")
                    root = ET.fromstring(content)
                    for elem in root.iter():
                        if elem.text and elem.text.strip():
                            text_parts.append(elem.text.strip())
                except Exception:
                    continue
        return "\n".join(text_parts)
    except Exception:
        return ""


def extract_from_pdf(file_bytes: bytes) -> str:
    """PDF에서 페이지별 텍스트 추출 (이미지 스캔본은 텍스트가 없을 수 있음)."""
    if not HAS_PYPDF:
        return ""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []
        for page in reader.pages:
            try:
                t = page.extract_text()
                if t:
                    parts.append(t)
            except Exception:
                continue
        return "\n".join(parts)
    except Exception:
        return ""


def extract_from_docx(file_bytes: bytes) -> str:
    """DOCX 문단 + 표 텍스트 추출."""
    if not HAS_DOCX:
        return ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception:
        return ""


def extract_from_xlsx(file_bytes: bytes) -> str:
    """XLSX(엑셀, ZIP+XML 구조)에서 셀 텍스트를 추출."""
    try:
        text_parts = []
        ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            shared_strings = []
            if "xl/sharedStrings.xml" in zf.namelist():
                content = zf.read("xl/sharedStrings.xml").decode("utf-8", errors="ignore")
                root = ET.fromstring(content)
                for si in root.findall(".//a:si", ns):
                    texts = [t.text or "" for t in si.findall(".//a:t", ns)]
                    shared_strings.append("".join(texts))

            sheet_files = sorted(
                n for n in zf.namelist()
                if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")
            )
            for sheet_name in sheet_files:
                content = zf.read(sheet_name).decode("utf-8", errors="ignore")
                root = ET.fromstring(content)
                for c in root.findall(".//a:c", ns):
                    v = c.find("a:v", ns)
                    if v is None or v.text is None:
                        continue
                    if c.get("t") == "s":
                        try:
                            text_parts.append(shared_strings[int(v.text)])
                        except (ValueError, IndexError):
                            pass
                    else:
                        text_parts.append(v.text)
        return "\n".join(text_parts)
    except Exception:
        return ""


def _select_best_from_zip(names):
    """ZIP 안에 여러 파일이 있으면 제안요청서/과업지시서 > hwp/hwpx > pdf/docx 순으로 하나 고름."""
    valid = [n for n in names if not n.endswith("/") and not os.path.basename(n).startswith(".")]
    for n in valid:
        base = os.path.basename(n)
        if "제안요청서" in base or "과업지시서" in base:
            return n
    for n in valid:
        if n.lower().endswith((".hwp", ".hwpx")):
            return n
    for n in valid:
        if n.lower().endswith((".pdf", ".docx")):
            return n
    return None


def extract_from_zip(file_bytes: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            best = _select_best_from_zip(zf.namelist())
            if not best:
                return ""
            inner_bytes = zf.read(best)
            return extract_text(inner_bytes, best)
    except Exception:
        return ""


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    파일 확장자에 따라 알맞은 추출 함수로 분기.
    확장자로 추출이 안 되거나(빈 결과), 확장자 자체가 불명확한 이름
    (예: API가 실제 파일명 대신 붙여준 '규격문서1' 같은 이름)인 경우,
    파일 내용의 시그니처(매직 바이트)를 보고 실제 형식을 다시 판별합니다.
    """
    ext = Path(filename).suffix.lower()
    text = ""
    if ext == ".hwp":
        text = extract_from_hwp(file_bytes)
    elif ext == ".hwpx":
        text = extract_from_hwpx(file_bytes)
    elif ext == ".pdf":
        text = extract_from_pdf(file_bytes)
    elif ext == ".docx":
        text = extract_from_docx(file_bytes)
    elif ext == ".xlsx":
        text = extract_from_xlsx(file_bytes)
    elif ext == ".zip":
        text = extract_from_zip(file_bytes)

    if text.strip():
        return text

    return extract_by_signature(file_bytes)


def extract_by_signature(file_bytes: bytes) -> str:
    """확장자 없이도, 파일 내용 맨 앞 바이트(시그니처)로 실제 형식을 추정해 추출."""
    if not file_bytes:
        return ""
    if file_bytes.startswith(b"%PDF"):
        return extract_from_pdf(file_bytes)
    if file_bytes.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"):
        return extract_from_hwp(file_bytes)  # 구버전 HWP/DOC/XLS는 모두 이 OLE 시그니처를 씁니다
    if file_bytes[:2] == b"PK":  # ZIP 기반 포맷(HWPX/DOCX/XLSX/일반 ZIP)
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                names = zf.namelist()
                if any(n.startswith("Contents/section") for n in names):
                    return extract_from_hwpx(file_bytes)
                if "word/document.xml" in names:
                    return extract_from_docx(file_bytes)
                if "xl/workbook.xml" in names:
                    return extract_from_xlsx(file_bytes)
        except Exception:
            pass
        return extract_from_zip(file_bytes)
    return ""


def download_and_extract(url: str, filename: str = "") -> str:
    """
    URL에서 파일을 내려받아 텍스트를 추출합니다.
    실패하면 빈 문자열을 반환합니다 (호출부에서 "추출 실패"로 간주해 금액 기준으로 대체 처리).
    """
    if not url:
        return ""
    if not filename:
        filename = os.path.basename(url.split("?")[0])
    try:
        resp = requests.get(url, timeout=30)
        resp.raise_for_status()
        text = extract_text(resp.content, filename)
        print(f"    [첨부파일 처리] {filename}: {len(resp.content)}바이트 다운로드, 텍스트 {len(text)}자 추출")
        return text
    except Exception as e:
        print(f"    [첨부파일 다운로드 실패] {filename}: {e}")
        return ""


def contains_keyword(text: str, keyword: str) -> bool:
    if not text:
        return False
    return keyword in text
